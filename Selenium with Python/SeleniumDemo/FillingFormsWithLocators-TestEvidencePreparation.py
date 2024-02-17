import time

import docx
import pyautogui
from docx import Document

from selenium import webdriver


# # from selenium.common import WebDriverException
# from selenium.webdriver.chrome.service import Service
# from selenium.webdriver.common.by import By
# from selenium.webdriver.support import expected_conditions
# from selenium.webdriver.support.select import Select
# from selenium.webdriver.support.wait import WebDriverWait


# Function to take screenshot and save it
def take_screenshot(driver, filename):
    driver.save_screenshot(filename)


# Function to scroll and capture screenshots
def scroll_and_capture(driver, num_screenshots):
    for i in range(num_screenshots):
        # Take screenshot
        filename = f"quote-section-{i+1}.png"
        take_screenshot(driver, filename)
        # Scroll down
        pyautogui.scroll(-1080)
        # Wait for page to load
        time.sleep(4)


# Function to paste screenshots into a Word document
def paste_into_word(screenshot_filenames, doc_filename):
    doc = Document()
    for filename in screenshot_filenames:
        doc.add_picture(filename, width=docx.shared.Inches(6))
        doc.add_paragraph()
    doc.save(doc_filename)


if __name__ == "__main__":
    # Initialize WebDriver
    driver = webdriver.Chrome()

    time.sleep(3)

    driver.get("https://rahulshettyacademy.com/angularpractice/")  # Replace with your URL

    time.sleep(10)

    # Scroll and capture screenshots
    num_screenshots = 2  # Change this to the number of screenshots you want to capture
    scroll_and_capture(driver, num_screenshots)

    time.sleep(2)

    # Close the WebDriver
    driver.quit()

    # Paste screenshots into Word document
    screenshot_filenames = [f"quote-section-{i+1}.png" for i in range(num_screenshots)]
    paste_into_word(screenshot_filenames, "test-evidence/webpage_screenshots.doc")
