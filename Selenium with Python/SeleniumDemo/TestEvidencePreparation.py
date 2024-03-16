import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx import Document

# Function to take screenshot and save it
def take_screenshot(driver, filename):
    driver.save_screenshot(filename)

# Function to scroll and capture screenshots
def scroll_and_capture(driver):
    # Get initial page height
    initial_height = driver.execute_script("return document.body.scrollHeight")

    # Start scrolling and capturing screenshots
    screenshot_filenames = []
    count = 1
    while True:
        filename = f"quote-section-{count}.png"
        take_screenshot(driver, filename)
        screenshot_filenames.append(filename)

        # Scroll down
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)  # Wait for page to load

        # Check if we've reached the end of the page
        new_height = driver.execute_script("return document.body.scrollHeight")
        if new_height == initial_height:
            break
        initial_height = new_height
        count += 1

    return screenshot_filenames

# Function to paste screenshots into a Word document
def paste_into_word(screenshot_filenames, doc_filename):
    doc = Document()
    for filename in screenshot_filenames:
        doc.add_picture(filename, width=doc.shared.Inches(6))
        doc.add_paragraph()
    doc.save(doc_filename)

if __name__ == "__main__":
    # Initialize WebDriver
    driver = webdriver.Chrome()
    driver.get("https://www.jetbrains.com/help/pycharm/set-up-a-git-repository.html M")  # Replace with your URL

    # Scroll and capture screenshots
    screenshot_filenames = scroll_and_capture(driver)

    # Close the WebDriver
    driver.quit()

    # Paste screenshots into Word document
    paste_into_word(screenshot_filenames, "webpage_screenshots.docx")
